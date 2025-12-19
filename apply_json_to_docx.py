import csv
import shutil
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------------------------------------------------
# CONFIG
# ----------------------------------------------------------------------
BOOKMARK_DOCX_DIR = Path("output/bookmarks")
TALLY_FILE = Path("output/bookmarks_tally.csv")
FINAL_TALLY_FILE = Path("output/final_tally.csv")

TOKEN_DOCX_DIR = Path("output/json")

DATA_DICT_FILE = Path("CWS-CARES Forms Data Dictionary-CountDataElements V0.01.xlsx")

# Cell text behavior (Word Table Properties > Cell > Options)
WRAP_TEXT = True                         # Wrap text
FIT_TEXT = False                         # Fit text (keep OFF)

# If True, also enforce constraints in headers/footers (optional)
INCLUDE_HEADERS_FOOTERS = False


# ----------------------------------------------------------------------
# DATA DICTIONARY (bookmark → token)
# ----------------------------------------------------------------------
def load_data_dictionary(path: Path) -> Dict[str, str]:
    df = pd.read_excel(path)

    if "Bookmark Name" not in df.columns:
        raise ValueError("data_dictionary.xlsx must contain 'Bookmark Name' column")
    if "Json Path" not in df.columns:
        raise ValueError("data_dictionary.xlsx must contain 'Json Path' column")

    mapping: Dict[str, str] = {}

    for _, row in df.iterrows():
        bookmark = row["Bookmark Name"]
        json_path = row["Json Path"]

        if pd.isna(bookmark) or pd.isna(json_path):
            continue

        bookmark_str = str(bookmark).strip()
        json_str = str(json_path).strip()
        if not bookmark_str or not json_str:
            continue

        token = json_str if (json_str.startswith("{{") and json_str.endswith("}}")) else f"{{{{{json_str}}}}}"
        mapping[bookmark_str] = token

    return mapping


# ----------------------------------------------------------------------
# TALLY CSV HANDLING
# ----------------------------------------------------------------------
def detect_filename_column(fieldnames: List[str]) -> str:
    candidates = ["filename", "file", "docx_file", "document", "doc_name", "document_name"]
    lower_map = {name.lower(): name for name in fieldnames}
    for cand in candidates:
        if cand in lower_map:
            return lower_map[cand]
    return fieldnames[0]


# ----------------------------------------------------------------------
# DOCX PARAGRAPH ITERATION (INCLUDES TABLES)
# ----------------------------------------------------------------------
def iter_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


# ----------------------------------------------------------------------
# TABLE ITERATION (BODY + NESTED)
# ----------------------------------------------------------------------
def iter_tables_from_tables(tables):
    for tbl in tables:
        yield tbl
        for row in tbl.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    yield from iter_tables_from_tables([nested])


def iter_tables(doc: Document):
    yield from iter_tables_from_tables(doc.tables)


def iter_all_tables(doc: Document):
    # Body
    yield from iter_tables(doc)

    if not INCLUDE_HEADERS_FOOTERS:
        return

    # Headers/Footers (optional)
    for section in doc.sections:
        yield from iter_tables_from_tables(section.header.tables)
        yield from iter_tables_from_tables(section.footer.tables)


# ----------------------------------------------------------------------
# TABLE CONSTRAINTS (PREVENT “EXPAND OUTWARD” WITHOUT MOVING HEADER TABLES)
# ----------------------------------------------------------------------
def set_table_fixed_layout(table):
    # Do NOT set alignment here (that’s what was moving header tables)
    table.autofit = False

    tblPr = table._tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _set_tblW(tblPr, width_twips: int):
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(width_twips))


def section_text_width_twips(doc: Document) -> int:
    # Use first section as reference
    s = doc.sections[0]
    # docx uses EMU for page dims/margins; 1 inch = 914400 EMU; 1 inch = 1440 twips
    width_emu = s.page_width - s.left_margin - s.right_margin
    width_inches = width_emu / 914400
    return int(round(width_inches * 1440))


def clamp_table_to_page_width(table, max_width_twips: int):
    tblPr = table._tbl.tblPr
    _set_tblW(tblPr, max_width_twips)


def get_tblgrid_widths_twips(table) -> List[int]:
    tblGrid = table._tbl.tblGrid
    if tblGrid is None:
        return []
    widths: List[int] = []
    for gc in tblGrid.findall(qn("w:gridCol")):
        w = gc.get(qn("w:w"))
        if w:
            widths.append(int(w))
    return widths


def set_cell_tcW_twips(cell, width_twips: int):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(width_twips))


def enforce_preserve_grid_widths(table):
    widths = get_tblgrid_widths_twips(table)
    if not widths:
        return
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths):
                break
            set_cell_tcW_twips(cell, widths[i])


def set_cell_wrap_and_fit(cell, wrap: bool = True, fit: bool = False):
    """
    Enforces:
      - Wrap text ON/OFF (w:noWrap absent/present)
      - Fit text ON/OFF (w:tcFitText present/absent)
    """
    tcPr = cell._tc.get_or_add_tcPr()

    noWrap = tcPr.find(qn("w:noWrap"))
    if wrap:
        if noWrap is not None:
            tcPr.remove(noWrap)
    else:
        if noWrap is None:
            tcPr.append(OxmlElement("w:noWrap"))

    fitText = tcPr.find(qn("w:tcFitText"))
    if fit:
        if fitText is None:
            tcPr.append(OxmlElement("w:tcFitText"))
    else:
        if fitText is not None:
            tcPr.remove(fitText)


def enforce_table_constraints(doc: Document, wrap_text: bool = WRAP_TEXT, fit_text: bool = FIT_TEXT):
    """
    Prevents runaway expansion while avoiding header-table “movement”:
      - fixed layout
      - clamp table width to text area
      - preserve existing tblGrid column widths (when present)
      - optional wrap/fit enforcement
    """
    max_twips = section_text_width_twips(doc)

    for table in iter_all_tables(doc):
        set_table_fixed_layout(table)
        clamp_table_to_page_width(table, max_twips)
        enforce_preserve_grid_widths(table)

        # Keep your cell behavior enforcement (safe)
        for row in table.rows:
            for cell in row.cells:
                set_cell_wrap_and_fit(cell, wrap=wrap_text, fit=fit_text)


# ----------------------------------------------------------------------
# APPLY DATA DICTIONARY TOKENS TO DOCX
# ----------------------------------------------------------------------
def apply_data_dict_to_doc(
    doc_path: Path, bookmark_to_token: Dict[str, str]
) -> Tuple[List[str], List[str]]:
    TOKEN_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(doc_path)

    # Constrain tables BEFORE replacement so Word can’t expand tables due to new text
    enforce_table_constraints(doc)

    bookmarks_replaced_set = set()
    json_used_set = set()

    for para in iter_paragraphs(doc):
        for run in para.runs:
            text = run.text or ""
            new_text = text

            for bookmark, token in bookmark_to_token.items():
                if bookmark in new_text:
                    new_text = new_text.replace(bookmark, token)
                    bookmarks_replaced_set.add(bookmark)
                    json_used_set.add(token)

            if new_text != text:
                run.text = new_text

    out_path = TOKEN_DOCX_DIR / doc_path.name
    doc.save(out_path)

    return sorted(bookmarks_replaced_set), sorted(json_used_set)


def copy_doc_without_changes(doc_path: Path) -> None:
    TOKEN_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out_path = TOKEN_DOCX_DIR / doc_path.name
    shutil.copy2(doc_path, out_path)


# ----------------------------------------------------------------------
# MAIN ORCHESTRATION
# ----------------------------------------------------------------------
def main() -> None:
    if not TALLY_FILE.exists():
        raise FileNotFoundError(f"Tally file not found: {TALLY_FILE}")
    if not DATA_DICT_FILE.exists():
        raise FileNotFoundError(f"Data dictionary not found: {DATA_DICT_FILE}")

    bookmark_to_token = load_data_dictionary(DATA_DICT_FILE)
    FINAL_TALLY_FILE.parent.mkdir(parents=True, exist_ok=True)

    with TALLY_FILE.open("r", encoding="utf-8", newline="") as f_in:
        reader = csv.DictReader(f_in)
        fieldnames = reader.fieldnames or []
        if not fieldnames:
            raise ValueError("Tally file has no columns")

        filename_col = detect_filename_column(fieldnames)
        out_fieldnames = fieldnames + ["bookmarks_replaced", "json"]

        with FINAL_TALLY_FILE.open("w", encoding="utf-8", newline="") as f_out:
            writer = csv.DictWriter(f_out, fieldnames=out_fieldnames)
            writer.writeheader()

            for row in reader:
                docx_name = row[filename_col]
                docx_path = BOOKMARK_DOCX_DIR / docx_name

                bookmarks_replaced: List[str] = []
                json_used: List[str] = []

                if docx_path.exists():
                    bookmarks_replaced, json_used = apply_data_dict_to_doc(
                        docx_path, bookmark_to_token
                    )

                row["bookmarks_replaced"] = ";".join(bookmarks_replaced)
                row["json"] = ";".join(json_used)
                writer.writerow(row)

    print(f"Wrote updated DOCX files to: {TOKEN_DOCX_DIR}")
    print(f"Wrote final tally CSV to: {FINAL_TALLY_FILE}")


if __name__ == "__main__":
    main()
