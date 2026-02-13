import csv
from docx import Document
from docx.shared import Inches

# Read the final_tally.csv
pairs = set()
with open('output/final_tally.csv', 'r', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        if row['bookmarks_replaced'] and row['json']:
            bookmarks = row['bookmarks_replaced'].split(';')
            jsons = row['json'].split(';')
            if len(bookmarks) == len(jsons):
                for b, j in zip(bookmarks, jsons):
                    pairs.add((b.strip(), j.strip()))

# Create the document
doc = Document()
doc.add_heading('Master Field Labels and Json Paths', 0)

# Add a table
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Field Label'
hdr_cells[1].text = 'Json Path'

for label, path in sorted(pairs):
    row_cells = table.add_row().cells
    row_cells[0].text = label
    row_cells[1].text = path

# Save the document
doc.save('output/converted/Main.docx')